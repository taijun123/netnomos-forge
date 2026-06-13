# -*- coding: utf-8 -*-
"""Build the NetNomos Forge product introduction DOCX.

Uses python-docx with explicit business-brief styling. The output is intended
for product walkthroughs and project handoff, not as source code documentation.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NetNomos-Forge_产品介绍.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 104, 118)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
GREEN_FILL = "EAF7EF"
RED_FILL = "FBEAEA"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    tbl.insert(0, grid)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def paragraph(doc, text="", style=None, before=None, after=None, align=None,
              size=11, color=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    color = BLUE if level in (1, 2) else DARK_BLUE
    size = 16 if level == 1 else 13 if level == 2 else 12
    set_run_font(run, size=size, color=color, bold=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    set_run_font(title_run, size=11, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    body_run = p2.add_run(body)
    set_run_font(body_run, size=10.5, color=RGBColor(35, 45, 55))
    paragraph(doc, "", after=4)


def add_kv_rows(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1800, 7560])
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        set_cell_shading(left, LIGHT_GRAY)
        for cell in (left, right):
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        r1 = left.paragraphs[0].add_run(label)
        set_run_font(r1, size=10.5, color=NAVY, bold=True)
        r2 = right.paragraphs[0].add_run(value)
        set_run_font(r2, size=10.5)
    paragraph(doc, "", after=4)


def add_matrix(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, header_fill)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        r = cell.paragraphs[0].add_run(header)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].paragraphs[0].paragraph_format.space_after = Pt(0)
            r = cells[j].paragraphs[0].add_run(value)
            set_run_font(r, size=10)
    set_table_geometry(table, widths)
    paragraph(doc, "", after=4)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    r = p.add_run("NetNomos Forge | Product Introduction")
    set_run_font(r, size=9, color=MUTED, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    r2 = fp.add_run("Internal demo brief - generated from project source")
    set_run_font(r2, size=9, color=MUTED)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    paragraph(doc, "产品介绍", before=8, after=0, size=11, color=MUTED, bold=True)
    title = paragraph(
        doc,
        "NetNomos Forge",
        before=0,
        after=4,
        size=28,
        color=NAVY,
        bold=True,
    )
    title.paragraph_format.line_spacing = 1.0
    paragraph(
        doc,
        "不改模型，只加规则：从规则自发现到受控生成的双轨合规工作台",
        after=14,
        size=13,
        color=MUTED,
        italic=True,
    )
    add_kv_rows(doc, [
        ("产品形态", "面向规则发现、资料核查、受控生成和合规报告的演示工作台"),
        ("核心技术", "NetNomos 规则挖掘，LeJIT 约束生成，RAG 规则解释，A/B 双轨报告"),
        ("当前阶段", "W4 稳定 demo：网络流量与财务报表两个垂直场景"),
        ("交付对象", "产品演示、技术评审、方案汇报和后续工程接力"),
    ])

    add_callout(
        doc,
        "一句话定位",
        "NetNomos Forge 将模型生成置于显式规则控制层之下，让用户看到规则从数据中被发现、被解释、被复用核查，并最终约束报告输出。",
    )

    heading(doc, "1. 产品价值", 1)
    paragraph(
        doc,
        "大模型可以生成流畅文本和结构化数据，但在财务、网络、安全、审计等场景中，流畅并不等于可信。NetNomos Forge 的产品假设是：模型外侧需要一层可审计的规则控制面，用显式规则约束关键数值、字段关系和业务逻辑。",
    )
    bullet(doc, "对业务用户：可以上传资料、输入报告问题，并看到规则命中、修正依据和 A/B 输出差异。")
    bullet(doc, "对技术团队：可以把 NetNomos 学出的规则转成统一 RuleSet/RuleCard，再交给后续验证、投影和生成环节。")
    bullet(doc, "对评审方：可以区分哪些规则来自数据自发现，哪些来自人工领域规则，避免把规则来源混为一谈。")

    doc.add_page_break()
    heading(doc, "2. 端到端工作流", 1)
    add_matrix(
        doc,
        ["阶段", "用户动作", "系统处理", "可见产物"],
        [
            ("资料接入", "上传或确认训练/待核查资料", "保存文件并登记 dataSourceId", "文件名、dataSourceId、上传状态"),
            ("规则学习", "启动规则学习或加载规则库", "NetNomos hitting-set/Z3 或 golden 规则加载", "规则数量、规则来源、事件流"),
            ("规则解释", "查看规则卡", "RAG 检索场景知识，可选 LLM 润色", "中文规则卡、citation、巧合提示"),
            ("资料核查", "上传新资料并运行核查", "RuleSet validate / Z3 check", "违规清单、命中规则、期望值"),
            ("双轨报告", "输入报告问题并运行 A/B", "A 轨裸模型；B 轨规则约束、投影和槽位回填", "标红报告、合规报告、干预日志"),
        ],
        [1500, 2200, 3200, 2460],
    )

    heading(doc, "3. 两个 W4 场景", 1)
    add_matrix(
        doc,
        ["场景", "演示资料", "规则重点", "A/B 对比重点"],
        [
            (
                "网络流量",
                "CIDDS NetFlow：正常 10k 训练流量 + 异常上传样本",
                "UDP Flags、Packets/Bytes 物理上下界、DNS 端口身份一致性",
                "A 轨生成/复述违规 NetFlow；B 轨输出 0 违规合规样本",
            ),
            (
                "财务报表",
                "960 行正确训练数据 + 华信咨询 8 期待审资料包",
                "进销存勾稽、资产负债配平、现金跨期、行业画像、比率背离",
                "A 轨照抄错误数值；B 轨修正营业成本等硬勾稽错误并给出风险提示",
            ),
        ],
        [1500, 2500, 2800, 2560],
    )

    heading(doc, "4. 系统架构", 1)
    paragraph(
        doc,
        "项目分为前端工作台、FastAPI 编排层、Forge Core SDK 和场景资产四层。前端负责交互和可视化，后端负责 job 编排和 SSE 事件，Forge Core 封装规则学习、解释、核查、投影和生成，场景目录保存数据规范、规则、知识库和演示样本。",
    )
    add_matrix(
        doc,
        ["层级", "模块", "职责"],
        [
            ("Web UI", "web/", "React/Vite 页面，展示规则卡、上传状态、进度条、核查表、A/B 报告和下载入口。"),
            ("Orchestrator", "server/", "FastAPI API、multipart 上传、后台 job、SSE 事件流、内存 store。"),
            ("Core SDK", "forge/core/", "规则引擎、RAG 解释器、LLM 路由、LeJIT 生成器、Projector、Reporter。"),
            ("Scenario Assets", "forge/scenarios/ 与 forge/rulesets/", "场景 spec、知识库、manual/golden 规则、LeJIT bundle、合规样本。"),
            ("Demo Assets", "demo_artifacts/w4_demo_assets/", "演示上传文件、正确训练数据、prompts、用户说明和限制说明。"),
        ],
        [1700, 2500, 5160],
    )

    heading(doc, "5. 当前实现状态与边界", 1)
    add_callout(
        doc,
        "已实现",
        "W4 已经实现页面文件选择、上传落盘、dataSourceId 登记、job 请求上下文透传、规则来源区分、手动核查触发、A/B 双轨问题输入、报告预览和 Markdown 下载。",
        GREEN_FILL,
    )
    add_callout(
        doc,
        "需要如实说明的边界",
        "当前 W4 上传资料用于触发流程和展示资料来源；核查与 A/B 结果仍复用稳定演示管线，尚未完成任意 CSV、PDF、Word 或 PCAP 文件的通用逐行解析。",
        RED_FILL,
    )
    bullet(doc, "Ollama 是可选增强；不可用时系统会降级到 codex 或 deterministic mock，保证演示可重复。")
    bullet(doc, "财务完整 workflow 可能需要几十秒，因为会运行完整校验、投影和报告生成管线。")
    bullet(doc, "当前 job store 是内存态，适合 demo；生产化需要持久化任务、资料和规则版本。")

    heading(doc, "6. 演示路径", 1)
    paragraph(doc, "演示入口说明见：demo_artifacts/w4_demo_assets/user.md。演示时只上传标注为上传文件的 CSV，其它正确数据、规则和样本用于讲解。")
    add_matrix(
        doc,
        ["Demo", "界面", "上传文件", "复制问题"],
        [
            (
                "财务",
                "http://127.0.0.1:5173/?v=w4source#/finance",
                "demo_artifacts/w4_demo_assets/finance/huaxin_audit_package.csv",
                "demo_artifacts/w4_demo_assets/finance/prompts.md",
            ),
            (
                "网络",
                "http://127.0.0.1:5173/?v=w4source#/network",
                "demo_artifacts/w4_demo_assets/network/netflow_rule_anomaly_upload.csv",
                "demo_artifacts/w4_demo_assets/network/prompts.md",
            ),
        ],
        [1200, 3050, 2850, 2260],
    )

    heading(doc, "7. 后续产品化方向", 1)
    numbered(doc, "接入真实文档解析：支持 PDF、Word、Excel、CSV、PCAP/NetFlow 的结构化抽取，并把上传资料逐行映射到核查结果。")
    numbered(doc, "规则版本治理：记录规则来源、训练数据、支持率、人工启停、版本差异和审批状态。")
    numbered(doc, "RAG 知识库管理：允许用户上传行业知识、制度文件和审计口径，并把引用透明展示在规则卡和报告中。")
    numbered(doc, "生产级任务系统：持久化 job、文件、规则集、报告和操作审计日志。")
    numbered(doc, "更多垂直场景：从网络和财务扩展到采购、合同、风控、资产运维和安全审计。")

    heading(doc, "8. 结论", 1)
    paragraph(
        doc,
        "NetNomos Forge 当前已经具备完整叙事闭环：从规则发现，到规则解释，到上传资料核查，再到 A/B 双轨报告。它展示的不是一个单点模型能力，而是一条面向高约束业务场景的模型控制路线：把规则作为可见、可解释、可复用的产品对象，放在生成流程之前和之中。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
